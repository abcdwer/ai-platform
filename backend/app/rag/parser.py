"""Document parser for various file formats."""
import io
import re
from typing import Optional, Dict, Any
from abc import ABC, abstractmethod
import httpx
from loguru import logger


class BaseParser(ABC):
    """Base document parser."""
    
    @abstractmethod
    def parse(self, content: bytes, **kwargs) -> str:
        """Parse document content and return text."""
        pass


class TextParser(BaseParser):
    """Plain text parser."""
    
    def parse(self, content: bytes, **kwargs) -> str:
        """Parse plain text."""
        return content.decode("utf-8", errors="ignore")


class MarkdownParser(BaseParser):
    """Markdown parser (strips markdown syntax)."""
    
    def parse(self, content: bytes, **kwargs) -> str:
        """Parse markdown and extract plain text."""
        text = content.decode("utf-8", errors="ignore")
        # Remove markdown links but keep text
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        # Remove image syntax
        text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', '', text)
        # Remove headers markers
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        # Remove bold/italic markers
        text = re.sub(r'\*{1,3}([^*]+)\*{1,3}', r'\1', text)
        text = re.sub(r'_{1,3}([^_]+)_{1,3}', r'\1', text)
        # Remove code blocks
        text = re.sub(r'```[\s\S]*?```', '', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # Remove blockquotes
        text = re.sub(r'^>\s+', '', text, flags=re.MULTILINE)
        # Remove horizontal rules
        text = re.sub(r'^[-*_]{3,}$', '', text, flags=re.MULTILINE)
        return text.strip()


class PDFParser(BaseParser):
    """PDF parser using PyPDF2."""
    
    def parse(self, content: bytes, **kwargs) -> str:
        """Parse PDF and extract text."""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            return "\n\n".join(text_parts)
        except ImportError:
            logger.error("PyPDF2 not installed, cannot parse PDF")
            raise ImportError("Please install PyPDF2: pip install PyPDF2")
        except Exception as e:
            logger.error(f"Failed to parse PDF: {e}")
            raise


class DocxParser(BaseParser):
    """Word document parser."""
    
    def parse(self, content: bytes, **kwargs) -> str:
        """Parse Word document."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            return "\n\n".join(paragraphs)
        except ImportError:
            logger.error("python-docx not installed, cannot parse DOCX")
            raise ImportError("Please install python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"Failed to parse DOCX: {e}")
            raise


class CSVParser(BaseParser):
    """CSV parser."""
    
    def parse(self, content: bytes, **kwargs) -> str:
        """Parse CSV and convert to text."""
        try:
            import csv
            import io as csv_io
            text_parts = []
            decoded = content.decode("utf-8", errors="ignore")
            reader = csv.reader(csv_io.StringIO(decoded))
            for row in reader:
                text_parts.append(" | ".join(row))
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to parse CSV: {e}")
            raise


class HTMLParser(BaseParser):
    """HTML parser for web content."""
    
    def parse(self, content: bytes, **kwargs) -> str:
        """Parse HTML and extract text."""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(content, "lxml")
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            # Get text
            text = soup.get_text()
            # Clean up whitespace
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            return "\n".join(chunk for chunk in chunks if chunk)
        except ImportError:
            logger.error("beautifulsoup4 or lxml not installed")
            raise ImportError("Please install beautifulsoup4 and lxml: pip install beautifulsoup4 lxml")
        except Exception as e:
            logger.error(f"Failed to parse HTML: {e}")
            raise


async def fetch_url(url: str) -> bytes:
    """Fetch content from URL."""
    async with httpx.AsyncClient(timeout=30.0) as client:
        response = await client.get(url)
        response.raise_for_status()
        return response.content


class URLParser(BaseParser):
    """URL content parser."""
    
    async def parse_async(self, url: str, **kwargs) -> Dict[str, Any]:
        """Fetch and parse URL content."""
        try:
            content = await fetch_url(url)
            content_type = self._detect_content_type(url, content)
            
            parser = get_parser(content_type)
            text = parser.parse(content)
            
            return {
                "content": text,
                "content_type": content_type,
                "url": url
            }
        except Exception as e:
            logger.error(f"Failed to fetch/parse URL {url}: {e}")
            raise
    
    def _detect_content_type(self, url: str, content: bytes) -> str:
        """Detect content type from URL or content."""
        url_lower = url.lower()
        if url_lower.endswith(".pdf"):
            return "pdf"
        elif url_lower.endswith((".doc", ".docx")):
            return "docx"
        elif url_lower.endswith(".csv"):
            return "csv"
        elif url_lower.endswith((".md", ".markdown")):
            return "md"
        elif url_lower.endswith((".html", ".htm")):
            return "html"
        else:
            # Try to detect from content
            try:
                text = content[:500].decode("utf-8", errors="ignore").strip()
                if text.startswith("<!DOCTYPE") or text.startswith("<html"):
                    return "html"
            except:
                pass
            return "html"  # Default to HTML


def get_parser(content_type: str) -> BaseParser:
    """Get parser for content type."""
    parsers = {
        "txt": TextParser(),
        "text": TextParser(),
        "md": MarkdownParser(),
        "markdown": MarkdownParser(),
        "pdf": PDFParser(),
        "docx": DocxParser(),
        "word": DocxParser(),
        "csv": CSVParser(),
        "html": HTMLParser(),
        "htm": HTMLParser(),
        "url": URLParser(),  # Special case
    }
    parser = parsers.get(content_type.lower())
    if parser is None:
        raise ValueError(f"Unsupported content type: {content_type}")
    return parser


async def parse_document(
    content: bytes,
    content_type: str,
    url: Optional[str] = None
) -> Dict[str, Any]:
    """Parse document content."""
    parser = get_parser(content_type)
    
    if isinstance(parser, URLParser):
        if url:
            return await parser.parse_async(url)
        else:
            raise ValueError("URL is required for URL content type")
    else:
        text = parser.parse(content)
        return {
            "content": text,
            "content_type": content_type,
            "url": url
        }
